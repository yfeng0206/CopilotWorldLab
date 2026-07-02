# -*- coding: utf-8 -*-
"""Copilot Automation Lab: final-pass proposal (clean arxiv-style Word doc, vision seat)."""
import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(_HERE, "CopilotAutomationLab_Proposal.docx")
FIGDIR = _HERE

# ----------------------------------------------------------------- references
REFS = [
 ("burger","Burger, B. et al. A mobile robotic chemist. Nature 583, 237-241 (2020)."),
 ("dai","Dai, T. et al. Autonomous mobile robots for exploratory synthetic chemistry. Nature 635, 890-897 (2024)."),
 ("crest","Zhang, Z. et al. A multimodal robotic platform for multi-element electrocatalyst discovery (CRESt). Nature 647, 390-396 (2025)."),
 ("vjepa2","Assran, M. et al. V-JEPA 2: self-supervised video models for understanding, prediction and planning. arXiv:2506.09985 (2025)."),
 ("vlajepa","Sun, J. et al. VLA-JEPA: enhancing vision-language-action models with a latent world model. arXiv:2602.10098 (2026)."),
 ("alab","Szymanski, N. J. et al. An autonomous laboratory for the accelerated synthesis of novel materials. Nature 624, 86-91 (2023)."),
 ("robochemist","RoboChemist: long-horizon and safety-compliant robotic chemical experimentation. CoRL 2025 (PMLR v305, 3537-3568). arXiv:2509.08820."),
 ("dreamtacvla","DreamTacVLA: a V-JEPA 2 world model for contact-rich coarse-to-fine insertion. arXiv:2512.23864 (2025)."),
 ("ahead","AHEAD: a latent world model that halts a policy when its predictive uncertainty crosses a threshold. arXiv:2606.02486 (2026)."),
 ("yu","Yu, C., Cai, Z., Pham, H. & Pham, Q.-C. Siamese convolutional neural network for sub-millimeter-accurate camera pose estimation and visual servoing. arXiv:1903.04713 (2019)."),
 ("labvla","LabVLA: grounding vision-language-action models in scientific laboratories. arXiv:2606.13578 (2026)."),
 ("tvfdit","TVF-DiT: real-laboratory imitation learning for test-tube and powder manipulation. arXiv:2603.01110 (2026)."),
 ("graspvla","Deng, S. et al. GraspVLA with SynGrasp-1B: a grasping foundation model for transparent-object grasping. arXiv:2505.03233 (2025)."),
 ("organa","Darvish, K. et al. ORGANA: a robotic assistant for automated chemistry. Matter 8(2), 101897 (2025). arXiv:2401.06949."),
 ("vlaperf","VLA-Perf: demystifying VLA inference performance (Jetson Thor latency benchmarks). arXiv:2602.18397 (2026)."),
 ("labutopia","Li, R. et al. LabUtopia: high-fidelity simulation and hierarchical benchmark for scientific embodied agents. NeurIPS 2025 (Datasets and Benchmarks). arXiv:2505.22634."),
 ("droid","Khazatsky, A. et al. DROID: a large-scale in-the-wild robot manipulation dataset. RSS 2024. arXiv:2403.12945."),
 ("oxe","O'Neill, A. et al. Open X-Embodiment: robotic learning datasets and RT-X. arXiv:2310.08864 (2023)."),
 ("robomind","Wu, K. et al. RoboMIND: multi-embodiment manipulation dataset (including UR-5e). arXiv:2412.13877."),
]
refnum = {k: i+1 for i,(k,_) in enumerate(REFS)}
def cite(*keys): return "[" + ", ".join(str(refnum[k]) for k in keys) + "]"

# ----------------------------------------------------------------- doc + helpers
doc = docx.Document()
SERIF = "Times New Roman"
doc.styles["Normal"].font.name = SERIF; doc.styles["Normal"].font.size = Pt(10.5)
for stylename in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    try:
        st = doc.styles[stylename]; st.font.name = SERIF
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except KeyError:
        pass

def para(text, sty=None, after=6, italic=False, size=None, align=None):
    p = doc.add_paragraph(style=sty); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    for i, seg in enumerate(text.split("**")):
        r = p.add_run(seg); r.bold = (i % 2 == 1)
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
    return p
def bullet(text): return para(text, sty="List Bullet", after=2)

_BM = [100]
def h(text, level, bm=None):
    p = doc.add_heading(text, level=level)
    if bm:
        _BM[0] += 1; bid = str(_BM[0])
        s = OxmlElement('w:bookmarkStart'); s.set(qn('w:id'), bid); s.set(qn('w:name'), bm)
        e = OxmlElement('w:bookmarkEnd'); e.set(qn('w:id'), bid)
        p._p.insert(0, s); p._p.append(e)
    return p
def link(text, anchor, indent=0.0, size=10):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    r.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    hl.append(r); p._p.append(hl); return p
def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""; run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""; run = cells[j].paragraphs[0].add_run(str(val)); run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths): row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t
def figure(name, caption, width=5.7):
    doc.add_picture(os.path.join(FIGDIR, name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after = Pt(9)
    if ". " in caption:
        label, rest = caption.split(". ", 1)
        rb = cp.add_run(label + ". "); rb.bold = True; rb.font.size = Pt(9); rb.font.color.rgb = RGBColor(0x22,0x22,0x22)
        rr = cp.add_run(rest); rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x22,0x22,0x22)
    else:
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x22,0x22,0x22)

CEN = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================ TITLE BLOCK
t = doc.add_heading("Learned World Model Manipulation for Self-Driving Chemistry Laboratories", level=0)
t.alignment = CEN
s = para("Copilot Automation Lab. The work pairs a \"grand plan\" with a first small scope demonstrator that "
         "realizes one component of that plan rather than the whole system.", after=4, align=CEN)
s.runs[0].italic = True
para("Goal: target workshop level artifacts for CVPR or NeurIPS 2027.", after=1, size=9, align=CEN)
para("Deadline: mid to late October 2026 for the demo POC (CVPR workshop).", after=1, size=9, align=CEN)
para("Date: 06/25/2026.", after=8, size=9, align=CEN)

# ----- required items
h("Required items and current availability", 2, bm="have")
tbl(["Item", "Role", "Status"],
 [["UR7e robotic arm", "manipulation", "available"],
  ["Robotiq 2F-85 / 2F-140 grippers", "grasping", "available"],
  ["Intel RealSense D405", "close range depth and RGB", "available"],
  ["NVIDIA Jetson Thor", "on robot inference", "available"],
  ["V-JEPA 2-AC, VLA-JEPA, GraspVLA checkpoints", "pretrained models", "public"],
  ["DROID, Open X-Embodiment, RoboMIND", "pretraining data", "public"],
  ["Training GPU (one A100 class)", "finetuning", "to arrange"],
  ["Teleoperation rig (UMI or GELLO)", "demonstrations", "to arrange"],
  ["Simulation", "demo POC and testing", "TBD"],
  ["Navigation stack (LiDAR-SLAM)", "moving between stations", "out of demo scope"]],
 widths=[2.4, 2.3, 1.5])

# ================================================================ CONTENTS
h("Contents", 2)
para("Click an entry to jump to that section.", after=4, italic=True, size=9)
for txt, anc in [("1. Overview", "s1"), ("2. Existing work and inspiration", "s2"),
                 ("3. Method", "s3"), ("    3.1 Navigation (classical)", "s31"),
                 ("    3.2 Planning and machine control (the Copilot)", "s32"),
                 ("    3.3 World model arm control", "s33"),
                 ("    3.4 Verification and guardrails", "s34"),
                 ("4. Action items", "s4"), ("5. Future work", "s5"), ("6. References", "s6")]:
    link(txt, anc, indent=0.25 if txt.startswith("    ") else 0.0)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ================================================================ 1 OVERVIEW
h("1. Overview", 1, bm="s1")
para(
 f"Automated and \"self-driving\" laboratories already operate as machines: they select, run and analyze "
 f"experiments with little human intervention {cite('burger','dai','alab','crest')}. In the systems deployed so "
 f"far, the decision making is automated, but the robotic arm is operated by classical, preprogrammed control: "
 f"scripted poses, per station touch point (\"cube\") calibration, and custom labware. Classical control is "
 f"precise and reliable, but each new instrument or labware item requires manual calibration, which is a "
 f"substantial part of the time and cost of deploying such a laboratory.")
para(
 f"We propose to add two learned components on top of a standard laboratory: (i) a language model (LLM) planner "
 f"(using Copilot) that turns a natural language request into a schedule of machine actions, and (ii) a world "
 f"model controller that performs the variable part of arm motion, namely the coarse approach to and placement of "
 f"labware at an instrument. The system keeps classical control for the precise and safeguard steps, while the "
 f"learned layer handles the variable approach and placement work that usually requires costly recalibration and "
 f"does not generalize well across new instruments.")
figure("fig_concept.png", "Figure 1. The laboratory is already a set of machines. We add a language model planner "
       "(Copilot) and a world model arm control layer, while instrument operation and navigation remain classical.")
para(
 f"The full system (the \"grand plan\") includes the real arm, navigation between stations, and real instruments. "
 f"As a first step we build a small scope software demo of the control loop, in which the arm and instruments are "
 f"simulated and the language model and world model are pluggable components. The demo covers LLM planning, "
 f"scheduling and the user interface; it does not include navigation, and does not claim real world manipulation "
 f"precision, which is addressed component by component in the action plan (Section 4).")
para("The work proceeds in three stages, of which the first is the immediate focus:")
bullet("**Stage 1 (current scope).** The software demo: the planner, scheduler and interface, the world model "
       "coarse placement, and the confidence gate measurement (Section 4.4), all evaluated in simulation. For the "
       "workshop, I believe simulation is sufficient at this stage, because the central question, whether the "
       "model's confidence predicts a failed handoff (classic to WM MPC), is answered over many simulated trials "
       "and does not need the physical arm.")
bullet("**Stage 2.** The same orchestrator on the physical UR7e for the single vial placement task (Section 4.1), "
       "which adds the real coarse placement error, the end to end seat success, and the setup cost comparison "
       "against classical calibration on real holders. (Cooper's lab took two years.)")
bullet("**Stage 3 (the grand plan).** Navigation between stations, real instruments, and further chemistry and "
       "labware (Section 5).")

# ================================================================ 2 EXISTING WORK
h("2. Existing work and inspiration", 1, bm="s2")
para(
 f"Two mobile robotic chemists {cite('burger','dai')} and several static self-driving laboratories "
 f"{cite('alab','crest')} demonstrate end to end autonomous experimentation. They differ in experiment setup and "
 f"scale but share an architecture: learned or automated experiment selection and classical, preprogrammed arm "
 f"control. To solve the fine positioning (millimeters), laser or touch feedback and a per station calibration "
 f"against a fixed reference are often deployed, with custom made grippers and racks. A separate, recent line of "
 f"work applies learned vision language action policies to laboratory manipulation {cite('labvla','robochemist','labutopia')}; "
 f"these show that learned arm control is feasible for laboratory tasks, but at coarse precision and mostly in "
 f"simulation or on low cost benchtop arms. The nearest system, RoboChemist (CoRL 2025) {cite('robochemist')}, "
 f"uses a single vision language model in three roles (planner, visual prompt generator, and success or compliance "
 f"monitor) together with a vision language action controller. Our design instead uses a video world model "
 f"(V-JEPA 2-AC or VLA-JEPA) as the controller, and a confidence signal from that same model (the predictive "
 f"energy of the V-JEPA 2-AC backend, or the disagreement across an ensemble for the VLA-JEPA backend) as a "
 f"candidate gate for the handoff to the precise seat: a signal from the model that generates the actions, rather "
 f"than a separate post hoc semantic monitor. Whether such a signal reliably marks the handoff is left open and "
 f"measured (Section 4.4).")
para(
 f"The closest learned manipulation precedents each cover one piece of this design but not their combination. "
 f"DreamTacVLA {cite('dreamtacvla')} uses a V-JEPA 2 world model for coarse to fine, contact rich insertion, but "
 f"end to end, with no classical fallback and no confidence gate. AHEAD {cite('ahead')} uses a world model's "
 f"predictive uncertainty as a threshold to halt a policy, but replans within that policy rather than handing off "
 f"to a classical controller. The main contribution is the specific combination that none of these makes: a "
 f"latent world model, instead of a vision language action policy, as the coarse controller, with a confidence "
 f"signal from that same model (its predictive energy or an ensemble disagreement) proposed as the handoff signal "
 f"to a visual servo submillimeter seat, and the reliability of that signal treated as the project's central open "
 f"question (Section 4.4).")
para("The capabilities have advanced over time, while a consistent limitation remains:")
tbl(["", "What improved", "What is still missing"],
 [["2020 " + cite('burger'), "A mobile arm serves many fixed instruments; closed loop optimization of one experiment type",
   "Per station cube calibration and custom labware; a single experiment type"],
  ["2024 " + cite('dai'), "Branching workflows over many chemistries; richer analysis (MS and NMR); modular instruments",
   "Arm control is still scripted and calibrated per station; no learned manipulation"],
  ["2025-26 " + cite('robochemist','labvla'), "Learned (VLA) arm control for lab tasks; language driven planning",
   "Coarse precision (no submillimeter); mostly simulation; not on a deployed mobile chemist"]],
 widths=[1.3, 3.0, 2.6])
para(
 f"Classical control is precise and deterministic, which is why deployed laboratories and factories rely on it, "
 f"and we do not propose reinventing the wheel. Its costs are the custom per station calibration and custom "
 f"labware it requires, and its limited ability to adapt to new instruments or changing conditions. Learned "
 f"policies adapt and generalize, but on their own they do not reach the precision that labware insertion needs "
 f"{cite('vjepa2')}; that precision is supplied by a deterministic visual servo {cite('yu')}. The opportunity is "
 f"to combine them: retain classical control for the precise motions, and use learning only for the variable "
 f"approach and placement and for translating human intent into machine plans. A latent world model paired with a "
 f"small action policy reaches useful manipulation from far fewer demonstrations than a vision language action "
 f"model trained from scratch: VLA-JEPA reports competitive results from about 100 demonstrations, one percent of "
 f"the data a conventional policy consumes {cite('vlajepa')}, which is the basis for expecting a lower per "
 f"laboratory data and setup cost. The learned layer has its own per deployment cost, so whether it lowers net "
 f"setup cost relative to per station calibration is an open question the project measures (Section 4.4). A "
 f"related benefit, operating outside a fixed precalibrated cell, depends on the navigation deferred to the grand "
 f"plan, and is therefore validated there rather than in the demonstrator.")

# ================================================================ 3 METHOD
h("3. Method", 1, bm="s3")
para(
 f"The system has three control layers plus navigation: (A) a language model planner and a deterministic "
 f"scheduler; (B) a world model controller for the variable arm motion; and (C) classical control, a visual "
 f"servo, for the precise motion. Navigation is classical and reused. The world model's own uncertainty "
 f"determines when control passes from (B) to (C).")
figure("fig_arch.png", "Figure 2. Control layers. The language model plans but does not issue real time motor "
       "commands; the scheduler dispatches; the world model performs the coarse motion; a visual servo performs "
       "the precise seat; an uncertainty gate decides the handoff. The planner and the scheduler together make up "
       "layer A.")

h("3.1 Navigation (classical)", 2, bm="s31")
para(
 f"Localization and movement between stations use classical LiDAR-SLAM and prebuilt station maps, as in the "
 f"mobile robotic chemists {cite('burger','dai')}. This component is reused unchanged and is not in the scope of "
 f"the demo, in which the arm is fixed base. It is included here only to place the rest of the system in context, "
 f"and is revisited in future work (Section 5).")

h("3.2 Planning and machine control (the Copilot)", 2, bm="s32")
para(
 f"A language model converts a natural language request into an ordered set of typed actions (a plan), and "
 f"updates the plan when results arrive. A deterministic scheduler executes the plan: it maintains a backlog, "
 f"resolves dependencies and the machine resources each action needs, dispatches actions that can run, and pushes "
 f"back items for those that cannot. This separation follows the control software of the 2024 mobile chemist "
 f"{cite('dai')}, with the rule based decision logic replaced by a language model and the scheduler kept "
 f"deterministic for reproducibility {cite('organa')}.")
bullet("**State.** Each machine has a status (idle, busy, error, or offline). Above the status is an inventory: "
       "every vial, sample and labware item maps to a location (a machine, a rack slot, or the gripper).")
bullet("**Actions.** A typed action records the machine resources it needs, its preconditions (over inventory and "
       "status), its effect, and whether it is run by classical or world model control.")
bullet("**Scheduling.** A transfer (moving a vial from machine X to machine Y) needs the arm and both X and Y; "
       "the scheduler acquires this whole set in one step or not at all, and requeues the action if any member is "
       "busy. Because no action ever holds one resource while waiting for another, the schedule cannot deadlock. "
       "When no action can run, the arm parks at an idle or charging position.")
para("In outline, the scheduler repeats:")
for ln in ["for each ready action whose resources are all free and preconditions hold:",
           "    lock its resources (atomically), dispatch it, schedule its completion",
           "if the arm is free and nothing can run: park it at the idle or charging position",
           "on completion: release resources, apply the effect, let the planner enqueue follow-ups"]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(ln); r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x20,0x20,0x20)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
para(
 f"The language model plans at a slow cadence and never issues real time motor commands; the scheduler and the "
 f"controllers run the real time loop. The demonstrator implements this layer with a simulated arm and a web "
 f"interface that shows the queue, machine status, inventory, and the actions as they run; the language model is a "
 f"pluggable component (a rule based planner in the demo, an API call in production).")

h("3.3 World model arm control", 2, bm="s33")
para(
 f"The world model controls the variable coarse approach and placement of labware, from the point at which a vial "
 f"is held to the point at which it is positioned just above the target holder. Two backends are available, and "
 f"they trade flexibility against speed:")
bullet("**Option 1, V-JEPA 2-AC with model predictive control** {0}. It plans to a goal image by minimizing a "
       "latent energy; this is flexible and needs no task specific demonstrations, but it plans slowly, on the "
       "order of tens of seconds per action (about 16 s per action was reported on a desktop GPU).".format(cite('vjepa2')))
bullet("**Option 2, VLA-JEPA feedforward policy** {0}. A fast policy that requires a small set of demonstrations, "
       "suitable for real time control on the on robot accelerator.".format(cite('vlajepa')))
para(
 "For the demo the primary path is Option 2 (the feedforward policy on the on robot accelerator), and Option 1 is "
 "the fallback when demonstrations cannot be collected. The slower planning latency of Option 1 is acceptable "
 "here because the arm motion is a single coarse step issued by the planner.")
para(
 f"The world model is not required to reach submillimeter precision. It places the vial within view of a wrist "
 f"camera; an image based visual servo then nulls the pixel error between the holder opening and the gripper axis "
 f"to the millimeter scale, and a passive chamfer completes the seat. Because the servo works in pixel space, it "
 f"is robust to the hand eye calibration error that causes the coarse miss, and it reaches the seat using only the "
 f"wrist camera, with no force, grip, or tactile sensor. A learned visual servo of this kind inserts connectors "
 f"with submillimeter clearance from vision alone, at about 0.6 mm and a 97.5% success rate {cite('yu')}. Public "
 f"checkpoints reach centimeter scale placement zero shot, so closing the remaining gap to the millimeter scale "
 f"with simulation and a few demonstrations is a measured outcome of the project, not an assumed capability; until "
 f"it is reached, a loosened holder clearance is used (Section 4.1).")
figure("fig_handoff.png", "Figure 3. The world model places the vial within view of the wrist camera; an image "
       "based visual servo nulls the pixel error to the millimeter scale, and a passive chamfer completes the "
       "seat. The precise motion uses vision only, with no force, grip, or tactile sensor, and the learned layer "
       "is not asked for submillimeter precision.")
para(
 f"The world model also provides a confidence estimate: the predictive energy of the V-JEPA 2-AC backend, or the "
 f"disagreement across an ensemble for the VLA-JEPA backend. The threshold is fit on a held out calibration set "
 f"to a target false accept rate; below it, the system performs a bounded retry and then flags a human, rather "
 f"than attempting the seat. Whether this estimate reliably predicts a failed handoff is an open question and the "
 f"project's central measurement (Section 4.4). The informative signal may be the model's own energy, the "
 f"disagreement across an ensemble, or the pixel error convergence of the visual servo, and the project compares "
 f"them against a simple baseline; if none improves on that baseline, the combine and gate approach does not hold.")

h("3.4 Verification and guardrails", 2, bm="s34")
para(
 f"The scheduler's preconditions and effects (Section 3.2) are not trusted on faith: a perception check verifies "
 f"them at run time, so the system acts on what it can observe, not only on its symbolic state. Before every "
 f"critical, hard to reverse step (releasing a vial into an instrument, inserting, or transporting), it confirms "
 f"the object is present and correctly held; after the step it confirms the intended effect (the vial is seated; "
 f"the gripper is now empty). On a mismatch it performs a bounded retry and then flags a human, and never proceeds "
 f"past a failed precondition. This is a guardrail independent of the learned policy, in the spirit of the monitor "
 f"and retry loop of RoboChemist {cite('robochemist')}.")
bullet("**Grasp check (classical, free).** The Robotiq gripper's commanded and actual finger width and grip force "
       "indicate whether a grasp succeeded, missed, or slipped: a deterministic precondition for any move, using "
       "the gripper's own feedback rather than a sensor added for the precise seat.")
bullet("**Scene check (two options).** Option 1: a vision language model inspects the RGB scene and returns "
       "success or a named condition (RoboChemist style " + cite('robochemist') + "), which is flexible and "
       "language driven but adds model latency. Option 2: a lightweight perception detector (open vocabulary "
       "detection, segmentation, or a pose check) verifies that the expected object is in the expected place, "
       "which is cheaper, deterministic and free of a model call.")
bullet("**Safety envelope.** A force torque hard limit and an emergency stop sit below the policy and the "
       "scheduler and are independent of them, as a safety limit rather than control; the demonstrator uses "
       "capped, inert surrogate liquids, so a missed grasp or a failed seat spills nothing.")

# ================================================================ 4 ACTION ITEMS
h("4. Action items", 1, bm="s4")
h("4.1 Demo target", 2)
para(
 f"The first target is a single task: the arm takes a vial from its onboard rack and places it into one "
 f"instrument holder to begin a short workflow. Vials are capped and contain liquid; no reagents are handled and "
 f"there is no navigation (a stationary arm for the demo). The task is first run as a software prototype (a "
 f"simulated arm and instruments with the planner and scheduler), then on the physical UR7e with the same "
 f"orchestrator. A generous holder clearance (a few millimeters) is used initially and tightened in later tests.")
para(
 f"The demo validates the control loop (planning, scheduling, the interface) and the confidence gate mechanism. "
 f"To exercise the learned layer despite the fixed base, the holder and base poses are perturbed synthetically to "
 f"stand in for the docking error a mobile platform would introduce, and at least two holder geometries are used. "
 f"A run is counted as a success if it reaches a target seating rate over a set number of trials at the loosened "
 f"clearance.")
h("4.2 Training and data", 2)
para(
 f"Only a few primitives are learned: the coarse approach and placement, and the grasp of transparent tubes, for "
 f"which a transparent object grasp model is a candidate {cite('graspvla')}. The known geometry grasp from the "
 f"onboard rack and the visual servo seat are deterministic and need no learned data. Pretraining uses public "
 f"models and datasets {cite('vjepa2','droid','oxe','robomind')}. No public dataset contains chemistry labware, "
 f"so labware specific data is generated in simulation and supplemented with a small number of teleoperated "
 f"demonstrations. This simulation and transfer path is the project's main risk, and the choice of simulator is "
 f"left open for now (TBD). VLA-JEPA reached competitive results with 100 demonstrations on real world pick and "
 f"place tasks {cite('vlajepa')}, roughly one percent of the data a conventional vision language action policy "
 f"consumes; chemistry labware is harder, hence the higher bound. A recent real laboratory imitation learning "
 f"study used about 400 demonstrations per task with no simulation or pretrained model, so the figure for this "
 f"hardware is to be measured {cite('tvfdit')}.")
h("4.3 Hardware and compute", 2)
para(
 f"The arm is a UR7e (pose repeatability ±0.03 mm), with Robotiq 2F-85 and 2F-140 grippers and an Intel RealSense "
 f"D405 (close range; ideal range about 7 to 50 cm, with submillimeter depth near 7 cm). The arm's mechanical "
 f"precision is well within the holder clearance, so the limiting factor is perception and control. Inference runs "
 f"on the on robot Jetson Thor: the feedforward policy (Section 3.3, Option 2) at interactive rates (a 2.7 billion "
 f"parameter feedforward policy is benchmarked at about 19 Hz on this device {cite('vlaperf')}), and the model "
 f"predictive backend (Option 1) as a slow, occasional step whose latency on Thor is to be measured (it is about "
 f"16 s per action on a desktop GPU, and Thor has lower memory bandwidth). An off board workstation GPU is an "
 f"alternative if the planning backend needs lower latency.")
h("4.4 Evaluation", 2)
para(
 f"Two measurements are primary: the confidence gate for the handoff, and setup cost. **(1) The confidence gate:** "
 f"whether the world model's confidence predicts a failed handoff, reported as area under the ROC curve against a "
 f"simple baseline and against the pixel error convergence of the visual servo; if it does not beat these, the "
 f"combine and gate approach does not hold. **(2) Setup cost:** the number of demonstrations (and wall clock time) "
 f"to reach a target seating rate on a new holder, against the time to calibrate the same holder by the classical "
 f"procedure, measured over at least two holders; the approach is worthwhile only if the per holder learned cost, "
 f"once one time pretraining and simulation are amortized over many stations, falls below the classical "
 f"calibration cost. Secondary measurements are the coarse placement error against the capture range, and the end "
 f"to end seating success with the bounded retry. All are reported with trial counts and confidence intervals, and "
 f"a negative result on either primary measurement is itself an informative finding.")

# ================================================================ 5 FUTURE
h("5. Future work", 1, bm="s5")
bullet("Navigation between stations (classical SLAM, integrated with the scheduler) to extend the demonstrator to "
       "a mobile platform.")
bullet("Operation with real reagents, subject to a safety review and a hazardous compound check on the planner.")
bullet("Additional chemistry and instruments, and measurement of generalization across labware without "
       "recalibration.")
bullet("On hardware latency measurements for the world model on the Jetson Thor.")

# ================================================================ 6 REFERENCES
h("6. References", 1, bm="s6")
for k, txt in REFS:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{refnum[k]}] "); r.bold = True; r.font.size = Pt(9)
    r2 = p.add_run(txt); r2.font.size = Pt(9)

doc.save(OUT)
print("SAVED", OUT)
print("refs", len(REFS))
